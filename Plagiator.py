import numpy as np
import docx
import os
import pathlib
import pandas as pd

from plotly.offline import plot
import plotly.graph_objects as go

from odf import text, teletype
from odf.opendocument import load

from nltk.tokenize import word_tokenize
import gensim

# path = pathlib.Path(__file__).parent.absolute()
path = 'C:\\Users\\joaco\\Desktop\\Joac\\scripts\\Plagio'
os.chdir(path)
os.chdir('.\\Files')

# DEFINO FUNCIONES
def sort_descending_tuple (a,b):
    if len(a)!= len(b): print('Error: Different lengths')
    for i in range(len(a)):
        indice = np.where(a[i:] == max(a[i:]))[0][0]+i
        valor_guardado = a[i]
        nombre_guardado = b[i]
        a[i] = a[indice]
        a[indice] = valor_guardado
        b[i] = b[indice]
        b[indice] = nombre_guardado

# CREO TABLA DE TEXTOS
print('\nLoading files...')
documents={}
Consignas = False
Error = False

for filename in os.listdir(os.getcwd()):
    print(filename)
    if filename[0] == '0':
        if str(filename)[-1] == 'x':
            Consignas_text = []
            allparas = docx.Document(str(filename)).paragraphs
            for i in range(len(allparas)):
                if allparas[i].text != '': Consignas_text.append(str(allparas[i].text))
            Consignas = True
        elif str(filename)[-1] == 't':
            Consignas_text = []
            doc = load(str(filename))
            allparas = doc.getElementsByType(text.P)
            for i in range(len(allparas)):
                if teletype.extractText(allparas[i]) != '': Consignas_text.append(str(teletype.extractText(allparas[i])))
    else:
        if str(filename)[-1] == 'x':
            documents[str(filename)] = []
            try:
                docx.Document(str(filename)).paragraphs
            except IOError:
                print ("\033[1;31;49m IOError: \033[1;37;49m Could not open file '{}'".format(str(filename)))
                Error = True
                pass
            except UnicodeDecodeError:
                print ("\033[1;31;49m UnicodeDecodeError: \033[1;37;49m Could not open file '{}'".format(str(filename)))
                Error = True
                pass
            allparas = docx.Document(str(filename)).paragraphs
            for i in range(len(allparas)):
                if allparas[i].text != '': documents[str(filename)].append(str(allparas[i].text))
        elif str(filename)[-1] == 't':
            try:
                doc = load(str(filename))
            except IOError:
                print ("\033[1;31;49m IOError: \033[1;37;49m Could not open file '{}'".format(str(filename)))
                Error = True
                pass
            except  UnicodeDecodeError:
                print ("\033[1;31;49m UnicodeDecodeError: \033[1;37;49m Could not open file '{}'".format(str(filename)))
                Error = True
                pass
            allparas = doc.getElementsByType(text.P)
            documents[str(filename)] = []
            for i in range(len(allparas)):
                if teletype.extractText(allparas[i]) != '': documents[str(filename)].append(str(teletype.extractText(allparas[i])))
        else:
            print("\033[1;31;49m Error: \033[1;37;49m Unknow format on file '{}'".format(str(filename)))
            Error = True

if Error: print('\nError files excluded from comparison')

# ARMADO DE CORPUS Y BOG ETC
print('\nMaking dictionary and corpus...')
gen_docs = []
gen_docs = [[w.lower() for w in word_tokenize(str(documents[str(filename)]))] for filename in documents.keys()]

dictionary = gensim.corpora.Dictionary(gen_docs)

corpus = [dictionary.doc2bow(gen_doc) for gen_doc in gen_docs]

tf_idf = gensim.models.TfidfModel(corpus)

sims = gensim.similarities.Similarity(os.getcwd(),tf_idf[corpus], num_features=len(dictionary))

# COMPARACION CON CONSIGNAS
print('\nComparing files...')

if Consignas:
    Con_Consignas = []
    Similitud_Consignas = []

    query_doc = [w.lower() for w in word_tokenize(str(Consignas_text))]
    query_doc_bow = dictionary.doc2bow(query_doc) #update an existing dictionary and create bag of words
    query_doc_tf_idf = tf_idf[query_doc_bow]
    results = sims[query_doc_tf_idf]
    nombres = [str(filename) for filename in documents.keys()]
    sort_descending_tuple(results, nombres)
    i=0
    while results[i] > 0.17:
        Con_Consignas.append(nombres[i])
        Similitud_Consignas.append('{:.1f}%'.format(results[i]*100))
        i += 1

# COMPARACION ENTRE ARCHIVOS CON CONSIGNAS
    archivo=[]
    Simil_to = []
    Similitud = []
    Sugerencias = []

    cantidad_por_file = 5

    for filename in documents.keys():
        query_doc = [w.lower() for w in word_tokenize(str(documents[str(filename)]))]
        query_doc_bow = dictionary.doc2bow(query_doc) #update an existing dictionary and create bag of words
        query_doc_tf_idf = tf_idf[query_doc_bow]
        results = sims[query_doc_tf_idf]
        nombres = [str(filename) for filename in documents.keys()]
        sort_descending_tuple(results, nombres)
        for i in range(cantidad_por_file):
            Ambos_Consignas = False
            archivo.append(str(filename))
            Simil_to.append(nombres[i+1])
            Similitud.append('{:.1f}%'.format(results[i+1]*100))
            if str(filename) in Con_Consignas and nombres[i+1] in Con_Consignas: Ambos_Consignas = True
            if results[i+1] < 0.4 and Ambos_Consignas: Sugerencias.append('Ambos con consignas')
            elif 0.4  < results[i+1] < 0.7 and Ambos_Consignas: Sugerencias.append('Palabras cambiadas. Ambos con consignas')
            elif 0.7  < results[i+1] < 1 and Ambos_Consignas: Sugerencias.append('Copia textual. Ambos con consignas')
            elif 0.3  < results[i+1] < 0.6 and Ambos_Consignas == False: Sugerencias.append('Palabras Cambiadas')
            elif 0.6  < results[i+1]  < 1 and Ambos_Consignas == False: Sugerencias.append('Copia textual')
            else: Sugerencias.append('-')

# COMPARACION ENTRE ARCHIVOS SIN DAR CONSIGNA
else:
    archivo=[]
    Simil_to = []
    Similitud = []
    Sugerencias = []
    cantidad_por_file = 6-1

    for filename in documents.keys():
        query_doc = [w.lower() for w in word_tokenize(str(documents[str(filename)]))]
        query_doc_bow = dictionary.doc2bow(query_doc) #update an existing dictionary and create bag of words
        query_doc_tf_idf = tf_idf[query_doc_bow]
        results = sims[query_doc_tf_idf]
        nombres = [str(filename) for filename in documents.keys()]
        sort_descending_tuple(results, nombres)
        for i in range(cantidad_por_file):
            archivo.append(str(filename))
            Simil_to.append(nombres[i+1])
            Similitud.append('{:.1f}%'.format(results[i+1]*100))
            if 0.3  < results[i+1] < 0.6: Sugerencias.append('Palabras cambiadas')
            elif 0.6  < results[i+1] < 1: Sugerencias.append('Copia textual')
            else: Sugerencias.append('-')

# ARMO DATAFRAME CON TABLA FINAL
print('\nPlotting results...')

tabla = {}
tabla['Archivo'] = archivo
tabla['Similar_a'] = Simil_to
tabla['Similitud'] = Similitud
tabla['Sugerencias'] = Sugerencias
df = pd.DataFrame(tabla, columns = ['Archivo', 'Similar_a', 'Similitud', 'Sugerencias'])


# PLOTEO TABLA FINAL
os.chdir(path)
rowEvenColor = 'white'
rowOddColor = 'lavender'
colores_ordenados = [rowOddColor]*cantidad_por_file+[rowEvenColor]*cantidad_por_file
fig = go.Figure(data=[go.Table(header=dict(values=list(df.columns), fill_color='paleturquoise',
                                           align='left'), cells=dict(values=[df.Archivo,
                                                       df.Similar_a, df.Similitud, df.Sugerencias],
                                                       fill_color=[colores_ordenados*int(len(archivo)/10)*len(tabla)],
                                                       align='left'))])
plot(fig, auto_open=True)

k=input("Press close to exit")


 

